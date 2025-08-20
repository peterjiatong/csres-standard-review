# util.py
from pathlib import Path
import sys
from datetime import datetime
from urllib.parse import quote_plus, urljoin
import logging
import requests
from bs4 import BeautifulSoup
import time
import re
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
import os
from dotenv import load_dotenv

load_dotenv()

# 寻找基础目录
def get_base_dir():
    """Get base directory for both frozen and non-frozen states"""
    if getattr(sys, 'frozen', False):
        return Path(sys.executable).parent
    else:
        return Path(__file__).resolve().parent

# 获取当前日期字符串
def get_current_date_string():
    """Get current date in MM_DD format"""
    return datetime.now().strftime("%m_%d")

# Initialize constants
BASE_DIR = get_base_dir()
MONTH_DAY = get_current_date_string()

# File paths
SRC_FILE = BASE_DIR / os.getenv("SRC")
DEST_FILE = BASE_DIR / os.getenv("DEST")

# Web scraping constants
BASE_URL = "http://www.csres.com/"
SEARCH_URL = urljoin(BASE_URL, "s.jsp")
ANTI_CRAWL_URL = "http://www.csres.com/error/noright.html"

# UA混淆
HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/125.0.0.0 Safari/537.36"),
    "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
               "image/avif,image/webp,image/apng,*/*;q=0.8,"
               "application/signed-exchange;v=b3;q=0.7"),
    "Accept-Language": "zh-CN,zh;q=0.9,ja-CN;q=0.8,ja;q=0.7,en-CN;q=0.6,en;q=0.5",
    "Accept-Encoding": "gzip, deflate",      # requests understands this
    "Referer": "http://www.csres.com/",
    "Connection": "keep-alive",
}

CURRENT = {"现行", "即将实施"}
# 中文标点符号到英文标点符号的映射
_ZH2EN = {
    "，": ",",  "。": ".",  "：": ":",  "；": ";",  "？": "?",
    "！": "!",  "（": "(",  "）": ")",  "【": "[",  "】": "]",
    "《": "<",  "》": ">",  "“": '"',  "”": '"',  "‘": "'", "’": "'",
    "、": ",",  "－": "-",  "～": "~", "＋": "+",  "％": "%",
}
# 编译正则表达式用于替换中文标点符号
_RE_ZH_PUNC = re.compile("|".join(map(re.escape, _ZH2EN)))

def zh_punc_to_en(text: str) -> str:
    # 把 text 中常见中文全角标点替换为英文半角标点
    return _RE_ZH_PUNC.sub(lambda m: _ZH2EN[m.group()], text)

# 正则表达式匹配标准代码
CODE_REGEX = r"""
(?P<code>                       # —— 整体命名捕获 "code"
    [A-Z]+                   # ① 主前缀  (GB / GBZ / YY123 …)
    ([0-9]+)?
    (?:/[A-Z0-9]+)?             # ② 可选子前缀 (/T / Z1 …)
    \s*                         # ③ 可选空格  (GB/T  5750)
    \d+(?:\.\d+)*               # ④ 数字主体 (5750  或  5750.1)
    (?:                         # ⑤ **可选** 数字范围
        [~\-～]                 #    连接符：~  -  或全角～  
        \d+(?:\.\d+)*           #    范围右端数字  (5750.13)
    )?                          #    ← 整个范围块可省略
    (?:-\d{4})?                 # ⑥ **可选** 年份后缀  -2019 (4位数字)
    (?:[A-Z]+)?                 # ⑦ **可选** 尾部字母  E / A / AB…
    (?:/XG\d+-\d{4})?           # ⑧ **可选** /XG 修订号 /XG1-2022
)
"""

# 基础：左右括号的可选包裹
P_LEFT  = r"[（(]"      # 全角（ 或 半角 (
P_RIGHT = r"[)）]"      # 半角 ) 或 全角 ）

PAT_CODE = re.compile(
    rf"""^
        ({P_LEFT}\s*)?        # 可选左括号
        {CODE_REGEX}          #   ← 原本的命名组 (?P<code>…)
        (\s*{P_RIGHT})?       # 可选右括号
        $""",
    re.VERBOSE
) 

PAT_NAME = re.compile(r"^《(?P<name>[^》]+)》$")   # 不变

PAT_LINE_1 = re.compile(
    rf"""{P_LEFT}?            # 可选左括号
        \s*
        {CODE_REGEX}          # ⇒ 仍然提供 (?P<code>)
        \s*
        {P_RIGHT}?            # 可选右括号
        \s*
        《(?P<name>[^》]+)》   #   名称
        .*?
    """,
    re.VERBOSE
)

PAT_LINE_2 = re.compile(
    rf"""《(?P<name>[^》]+)》  # 名称
        .*?
        \s*
        {P_LEFT}?             # 左括号可选
        \s*
        {CODE_REGEX}          # ⇒ (?P<code>)
        \s*
        {P_RIGHT}?            # 右括号可选
    """,
    re.VERBOSE
)

def update_std_index(df_has_output):
    df_std = (
        df_has_output[["标准编号", "标准名称", "状态", "替代情况"]]
        .astype(str)
        .apply(lambda s: s.str.strip())
    )
    df_std["标准编号"] = df_std["标准编号"].str.replace(r"\s+", "", regex=True)
    df_std["标准名称"] = df_std["标准名称"].apply(zh_punc_to_en)
    return {row["标准编号"]: (row["状态"], row["标准名称"], row["替代情况"]) for _, row in df_std.iterrows()}

# 根据提供的路径和文件名生成唯一的日志文件路径
def get_path_for_log_file(path, file_name):
    """Generate unique log file path with date and index"""
    output_dir = BASE_DIR / Path(path)
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = Path(file_name).stem
    suffix = Path(file_name).suffix
    glob_pat = f"{stem}_{MONTH_DAY}_*{suffix}"
    existing = sorted(output_dir.glob(glob_pat))
    next_idx = len(existing) + 1
    new_name = f"{stem}_{MONTH_DAY}_{next_idx}{suffix}"
    return output_dir / new_name

# 根据提供的路径和文件名生成唯一的报告文件夹路径，并创建报告文件
def get_path_for_report_folder(path_stem, file_name):
    """Generate unique report folder path and create folder"""
    parent_dir = BASE_DIR
    glob_pat = f"{path_stem}_{MONTH_DAY}_*"
    existing = sorted(d for d in parent_dir.glob(glob_pat) if d.is_dir())
    next_idx = len(existing) + 1
    new_dir = parent_dir / f"{path_stem}_{MONTH_DAY}_{next_idx}"
    new_dir.mkdir(parents=True, exist_ok=True)
    return new_dir / file_name

# 自定义异常类
class CrawlError(Exception):
    """Custom exception for web crawling errors"""
    def __init__(self, msg, code, req_headers, resp_headers):
        super().__init__(msg)
        self.code = code
        self.req_headers = req_headers
        self.resp_headers = resp_headers

# 找到表格单元格后提取文本
def _text_after(label, soup):
    """Extract text from table cell after finding label"""
    if label == "替代情况":
        # Special handling for 替代情况
        # Look for any td containing "替代情况" (regardless of images)
        for td in soup.find_all("td"):
            if td.get_text(strip=True).startswith("替代情况"):
                next_td = td.find_next_sibling("td")
                if next_td:
                    return next_td.get_text(strip=True)
        return None
    else:
        # Original approach for other fields
        td = soup.find("td", string=re.compile(label))
        if td and td.find_next_sibling("td"):
            return td.find_next_sibling("td").get_text(strip=True)
        return None

# 设置cookie jar
def get_jar():
    """Create cookie jar for website authentication"""
    jar = requests.cookies.RequestsCookieJar()
    jar.set("source", "www.csres.com")
    jar.set("userName", f'"{os.getenv("CSRES_USERNAME")}"')
    jar.set("userPass", os.getenv("CSRES_PASSWORD"))
    logging.info("已设置cookie jar")
    return jar

# 创建搜索URL，支持中文GBK编码
def _search_url_gbk(keyword: str, page: int = 1) -> str:
    """Create search URL with GBK encoding for Chinese characters"""
    kw_gbk = quote_plus(keyword, encoding="gbk")
    return f"{SEARCH_URL}?keyword={kw_gbk}&pageNum={page}"

# 爬取单个标准代码的数据
def crawl_one_code(code, session, is_wrong_before=False):
    """Crawl data for a single standard code"""
    url_gbk = _search_url_gbk(code, page=1)

    # Retry up to 5 times for anti-crawl protection
    r_attempt = 0
    while r_attempt < 5:
        try:
            r = session.get(url_gbk, headers=HEADERS, timeout=(5, 30))
            soup = BeautifulSoup(r.text, "lxml")
            rows = soup.select('table.heng tr[bgcolor="#FFFFFF"]')
            
            # Handle known bad codes differently
            if is_wrong_before:
                if r.url != ANTI_CRAWL_URL and not rows:
                    raise CrawlError("无搜索结果", code, r.request.headers, r.headers)
            
            # Check for anti-crawl page and search results
            if r.url != ANTI_CRAWL_URL and rows:
                break
                
        except requests.exceptions.RequestException as e:
            logging.warning(f"网络请求失败 (尝试 {r_attempt + 1}/5): {e}")
        
        r_attempt += 1
        time.sleep(2)
    
    # Handle failure after all attempts
    if r_attempt == 5:
        if r.url == ANTI_CRAWL_URL:
            raise CrawlError("网站拒绝我们访问（www.csres.com/error/noright.html）", 
                           code, r.request.headers, r.headers)
        raise CrawlError("无搜索结果", code, r.request.headers, r.headers)

    # Validate search results
    if not rows:
        raise CrawlError("无搜索结果", code, r.request.headers, r.headers)
    elif len(rows) > 20:
        raise CrawlError("搜索结果过多（大于20个），请检查", code, r.request.headers, r.headers)

    # Process each result row
    hits = []
    for row in rows:
        try:
            href = row.find("a")["href"]
            status = row.find_all("td")[-1].get_text(strip=True)
            std_name = row.find_all("td")[1].get_text(strip=True)
            std_code = row.find_all("td")[0].get_text(strip=True)

            # Get detailed information from sub-page
            r2_attempt = 0
            while r2_attempt < 5:
                try:
                    r2 = session.get(urljoin(BASE_URL, href), headers=HEADERS, timeout=(5, 30))
                    soup2 = BeautifulSoup(r2.text, "lxml")
                    
                    if (r2.url != ANTI_CRAWL_URL and 
                        (_text_after("发布日期", soup2) or _text_after("实施日期", soup2) or _text_after("作废日期", soup2))):
                        break
                        
                except requests.exceptions.RequestException as e:
                    logging.warning(f"子页面请求失败 (尝试 {r2_attempt + 1}/5): {e}")
                
                r2_attempt += 1
                time.sleep(2)
        
            if r2_attempt == 5 and r2.url == ANTI_CRAWL_URL:
                raise CrawlError("子页面拒绝我们访问（www.csres.com/error/noright.html）", 
                               code, r2.request.headers, r2.headers)

            # Extract detailed information
            replacement_info = ""
            if status in ("作废", "废止"):
                replacement_info = _text_after("替代情况", soup2) or ""
                # print(replacement_info)
            
            info = {
                "标准编号": std_code,
                "标准名称": std_name,
                "状态": status,
                "发布日期": _text_after("发布日期", soup2) or "",
                "实施日期": _text_after("实施日期", soup2) or "",
                "作废日期": _text_after("作废日期", soup2) if status in ("作废", "废止") else "",
                "替代情况": replacement_info,
            }
            hits.append((info, r2.text))
            
        except Exception as e:
            logging.warning(f"处理行数据时出错: {e}")
            continue

    return hits

# 处理多个标准代码的爬取和数据存储
def process_code(code, session, df_has_output, df_no_output_or_too_much_outputs, 
                df_date_empty, df_err, is_wrong_before=False, max_retry=9, retry_sleep=5.0):
    """Process a single code with retry logic"""
    
    for attempt in range(max_retry + 1):
        try:
            hits = crawl_one_code(code, session, is_wrong_before)
            
            for info, r2text in hits:
                # Add to results dataframe
                df_has_output.loc[len(df_has_output)] = {
                    "标准编号": info["标准编号"],
                    "标准名称": info["标准名称"],
                    "状态": info["状态"],
                    "发布日期": info["发布日期"],
                    "实施日期": info["实施日期"],
                    "作废日期": info["作废日期"],
                    "替代情况": info["替代情况"],
                    "结果添加日期": MONTH_DAY,
                }

                # Check for missing dates (debug purposes)
                if (info["发布日期"] == "" and info["实施日期"] == "" and info["作废日期"] == ""):
                    logging.warning(f"⚠️  {code}: 可能没有发布日期、实施日期或作废日期")
                    df_date_empty.loc[len(df_date_empty)] = [info["标准编号"], r2text, MONTH_DAY]

            logging.debug(f"✅  {code}: 共处理{len(hits)}个结果")
            return

        except CrawlError as ce:
            # Handle known crawl errors
            if str(ce) in ["无搜索结果", "搜索结果过多（大于20个），请检查"]:
                df_no_output_or_too_much_outputs.loc[len(df_no_output_or_too_much_outputs)] = {
                    "标准编号": ce.code,
                    "错误信息": str(ce),
                    "结果添加日期": MONTH_DAY,
                }
                logging.warning(f"❌  {code}: {ce}")
                return
            else:
                df_err.loc[len(df_err)] = {
                    "标准编号": ce.code,
                    "错误信息": str(ce),
                    "Request-Headers": dict(ce.req_headers),
                    "Response-Headers": dict(ce.resp_headers),
                    "结果添加日期": MONTH_DAY,
                }
                logging.error(f"❌  {code}: {ce}")
                return

        except Exception as e:
            # Handle unexpected errors with retry
            if attempt < max_retry:
                logging.warning(f"⚠️  {code}: 尝试 {attempt + 1}/{max_retry + 1} 失败: {e}")
                time.sleep(retry_sleep)
                continue
            else:
                # Final failure
                df_err.loc[len(df_err)] = {
                    "标准编号": code,
                    "错误信息": f"{e}",
                    "Request-Headers": {},
                    "Response-Headers": {},
                    "结果添加日期": MONTH_DAY,
                }
                logging.error(f"❌  {code}: {e}，尝试{max_retry + 1}次仍失败")
                return

# 配置日志记录
def setup_logging(file_name):
    """Setup logging configuration"""
    log_path = get_path_for_log_file("log", file_name)
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        filename=log_path,
        filemode="w",
    )

# 加载现有数据
def load_existing_data():
    try:
        code_ok = pd.read_excel(SRC_FILE, sheet_name="有搜索结果的标准")["标准编号"].tolist()
        code_err = pd.read_excel(SRC_FILE, sheet_name="无搜索结果或搜索结果过多的标准")["标准编号"].tolist()
        known_codes = set(code_ok)
        # Clean whitespace from codes
        code_ok = [str(c).replace(" ", "") for c in code_ok]
        code_err = [str(c).replace(" ", "") for c in code_err]

        logging.info(f"原excel中“有搜索结果的标准”表长度为:{len(code_ok)}, “无搜索结果或搜索结果过多的标准”表长度为:{len(code_err)}")
        print(f"原excel中“有搜索结果的标准”表长度为:{len(code_ok)}, “无搜索结果或搜索结果过多的标准”表长度为:{len(code_err)}")

        return code_ok, code_err, known_codes
    
    except FileNotFoundError:
        logging.error(f"找不到源文件: {SRC_FILE}")
        raise
    except Exception as e:
        logging.error(f"读取Excel文件时出错: {e}")
        raise

# 初始化所有需要的DataFrame
def initialize_dataframes():
    """Initialize all required DataFrames"""
    df_has_output = pd.DataFrame(columns=[
        "标准编号", "标准名称", "状态", "发布日期", "实施日期", "作废日期", "替代情况", "结果添加日期"
    ])
    df_no_output_or_too_much_outputs = pd.DataFrame(columns=[
        "标准编号", "错误信息", "结果添加日期"
    ])
    df_date_empty = pd.DataFrame(columns=[
        "标准编号", "r2.text", "结果添加日期"
    ])
    df_err = pd.DataFrame(columns=[
        "标准编号", "错误信息", "Request-Headers", "Response-Headers", "结果添加日期"
    ])

    logging.info("初始化df")
    
    return df_has_output, df_no_output_or_too_much_outputs, df_date_empty, df_err

# 去重
def remove_duplicates(df_has_output, df_no_output_or_too_much_outputs, df_date_empty, df_err):
    """Remove duplicates from all dataframes"""    
    df_has_output = df_has_output.drop_duplicates(subset="标准编号", keep="first").reset_index(drop=True)
    df_no_output_or_too_much_outputs = df_no_output_or_too_much_outputs.drop_duplicates(subset="标准编号", keep="first").reset_index(drop=True)
    df_date_empty = df_date_empty.drop_duplicates(subset="标准编号", keep="first").reset_index(drop=True)
    df_err = df_err.drop_duplicates(subset="标准编号", keep="first").reset_index(drop=True)
    logging.info("已去重")
    
    return df_has_output, df_no_output_or_too_much_outputs, df_date_empty, df_err

# 保存Excel文件并调整格式
def save_excel_with_formatting(file_path, df_has_output, df_no_output_or_too_much_outputs, df_date_empty, df_err):
    # Save to Excel
    with pd.ExcelWriter(file_path, engine="openpyxl") as writer:
        df_has_output.to_excel(writer, sheet_name="有搜索结果的标准", index=False)
        df_no_output_or_too_much_outputs.to_excel(writer, sheet_name="无搜索结果或搜索结果过多的标准", index=False)
        df_err.to_excel(writer, sheet_name="报错(debug用)", index=False)
        df_date_empty.to_excel(writer, sheet_name="标准无详细日期(debug用)", index=False)
    
    # Adjust column widths
    wb = load_workbook(file_path)
    sheet_names = ["有搜索结果的标准", "无搜索结果或搜索结果过多的标准", "报错(debug用)", "标准无详细日期(debug用)"]
    
    for sheet_name in sheet_names:
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for col_idx, col in enumerate(ws.iter_cols(values_only=True), 1):
                max_len = max((len(str(v)) for v in col if v is not None), default=0)
                ws.column_dimensions[get_column_letter(col_idx)].width = max(max_len + 2, 10)
    
    wb.save(file_path)

# 生成新增标准报告
def generate_new_standards_report(df_has_output, known_codes):
    """Generate report for new standards"""
    df_new = df_has_output[~df_has_output["标准编号"].isin(known_codes)]
    
    if df_new.empty:
        logging.info("本次运行没有新增的标准记录")
        return
    
    txt_path = get_path_for_report_folder("更新数据库.exe的运行结果", "标准更新报告.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(f"新增标准统计（{MONTH_DAY}） 共 {len(df_new)} 条\n")
        f.write("-" * 78 + "\n")
        
        for _, row in df_new.iterrows():
            line = (
                f"{row['标准编号']:<20}"
                f"{row['状态']:<6}"
                f"标准名称: {row['标准名称']}"
            )
            f.write(line + "\n")
    
    logging.info(f"📝 已生成新增标准 TXT 报告: {txt_path}")

def generate_new_standards_report_in_exist_folder(folder, df_has_output, known_codes):
    """Generate report for new standards"""
    df_new = df_has_output[~df_has_output["标准编号"].isin(known_codes)]
    
    if df_new.empty:
        logging.info("本次运行没有新增的标准记录")
        return
    
    txt_path = folder / f"标准更新报告.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(f"新增标准统计（{MONTH_DAY}） 共 {len(df_new)} 条\n")
        f.write("-" * 78 + "\n")
        
        for _, row in df_new.iterrows():
            line = (
                f"{row['标准编号']:<20}"
                f"{row['状态']:<6}"
                f"标准名称: {row['标准名称']}"
            )
            f.write(line + "\n")
    
    logging.info(f"📝 已生成新增标准 TXT 报告: {txt_path}")

# Additional filter function to exclude Chinese-style codes
def is_valid_standard_code(code_text):
    """
    Filter out Chinese administrative codes and only keep alphanumeric standard codes
    """
    # Remove whitespace for checking
    clean_code = code_text.replace(" ", "")
    
    # Exclude patterns that contain Chinese characters or administrative terms
    chinese_patterns = [
        r'[\u4e00-\u9fff]',  # Any Chinese characters
        r'主席令',            # Presidential decree
        r'国务院令',          # State Council decree  
        r'号',               # Number (Chinese)
        r'第\d+号',          # "No. X" pattern
        r'〔\d+〕',          # Bracket notation like 〔2024〕
    ]
    
    for pattern in chinese_patterns:
        if re.search(pattern, code_text):
            return False
    
    # Must start with letters (standard prefixes)
    if not re.match(r'^[A-Z]+', clean_code):
        return False
        
    # Must contain numbers (all standards have numbers)
    if not re.search(r'\d', clean_code):
        return False
    
    return True

def extract_from_docx(path: Path):
    doc = Document(path)
    results = []

    def feed_text(text: str):
        # 去除空格
        text = text.strip()
        if not text:
            return False
        found = False        
        for pat in (PAT_LINE_1, PAT_LINE_2):
            for m in pat.finditer(text):
                orig_code = m.group("code")
                if not is_valid_standard_code(orig_code):
                    continue
                code = orig_code.replace(" ", "")
                name = zh_punc_to_en(m.group("name").strip())
                results.append((orig_code, code, name))
                found = True
        return found

    # 段落
    prev = ""
    check = False
    for i in doc.paragraphs:
        p = i.text.replace('\r','\n').split('\n')
        for line in p:
            cur = line
            if check:
                feed_text(prev + cur)
                check = False
                continue
            found = feed_text(line)
            if not found:
                prev = cur               # shift look-back
                check = True
    # 表格
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if not cells:
                continue

            # 行级“一对”匹配
            codes = [c for c in cells if PAT_CODE.fullmatch(c.replace(" ", ""))]
            names = [m.group("name") for m in map(PAT_NAME.fullmatch, cells) if m]
            if len(codes) == 1 and len(names) == 1:
                results.append((codes[0], codes[0].replace(" ", ""), zh_punc_to_en(names[0])))
                continue

            # 否则逐行逐格匹配
            for cell in cells:
                for line in cell.splitlines():
                    feed_text(line)
    return results

def normalize_name(s: str) -> str:
    """比较前，对名称再做一次统一：去空格 + 半角化"""
    return zh_punc_to_en(s).replace(" ", "")
